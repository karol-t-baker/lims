"""Certificate PDF generation — parse docx templates, fill results, render via weasyprint."""
import os
import re
from datetime import date, datetime
from pathlib import Path

from flask import render_template

from mbr.cert_mappings import CERT_MAPPINGS

SWIADECTWA_DIR = Path(__file__).parent.parent / "docs" / "swiadectwa"
WZORY_DIR = SWIADECTWA_DIR / "wzory"
OUTPUT_DIR = Path(__file__).parent.parent / "data" / "swiadectwa"

# Short product name tokens to match filenames
_PRODUCT_TOKENS = {
    "Chegina_K40GLOL": ["K40GLOL", "GLOL40", "GLOL"],
    "Chegina_K40GLO": ["K40GLO", "K40 GLO"],
    "Chegina_K40GL": ["K40GL", "K40 GL"],
    "Chegina_K7": ["K7"],
}

# For display: strip leading common prefixes
_DISPLAY_STRIP_RE = re.compile(
    r"^(Świadectwo_Certificate[-\s]+|Świadectwo-Certificate[-\s]+|"
    r"ŚwiadectwoCertificate[-\s]+|Certificate Świadectwo[-\s]+|"
    r"AVON\s+|LEHVOSS\s+|PRIME\s+|REVADA\s+)",
    re.IGNORECASE,
)


def _short_name_tokens(produkt: str) -> list[str]:
    """Return list of string tokens that identify this product in filenames."""
    if produkt in _PRODUCT_TOKENS:
        return _PRODUCT_TOKENS[produkt]
    # Fallback: try last part of product name e.g. "K7" from "Chegina_K7"
    parts = produkt.replace("Chegina_", "").split("_")
    return parts if parts else [produkt]


def _matches_product(filename: str, produkt: str) -> bool:
    fn_upper = filename.upper()
    for token in _short_name_tokens(produkt):
        # Match the token surrounded by non-alpha boundaries to avoid K40GL matching K40GLO
        pattern = re.escape(token.upper().replace(" ", r"\s*"))
        if re.search(r"(?<![A-Z])" + pattern + r"(?![A-Z])", fn_upper):
            return True
    return False


def _display_name(filename: str) -> str:
    """Remove common prefix and .docx suffix for human-readable name."""
    name = filename
    if name.lower().endswith(".docx"):
        name = name[:-5]
    name = _DISPLAY_STRIP_RE.sub("", name).strip()
    return name


def list_templates_for_product(produkt: str) -> list[dict]:
    """Return list of {filename, display} dicts for .docx templates matching produkt."""
    results = []
    seen = set()

    for search_dir in [SWIADECTWA_DIR, WZORY_DIR]:
        if not search_dir.exists():
            continue
        for path in sorted(search_dir.glob("*.docx")):
            fn = path.name
            if fn in seen:
                continue
            if _matches_product(fn, produkt):
                seen.add(fn)
                results.append({
                    "filename": fn,
                    "display": _display_name(fn),
                    "path": str(path),
                })

    results.sort(key=lambda x: x["display"].lower())
    return results


def _find_template_path(filename: str) -> Path:
    """Locate docx file in swiadectwa dirs."""
    for search_dir in [SWIADECTWA_DIR, WZORY_DIR]:
        p = search_dir / filename
        if p.exists():
            return p
    raise FileNotFoundError(f"Template not found: {filename}")


def parse_docx_template(filename: str) -> dict:
    """Parse a .docx certificate template.

    Returns:
        {
            "meta": {
                "tds": str,
                "cas": str,
                "opinion": str,
                "produkt_name": str,
            },
            "header": [str, ...],
            "rows": [{"param": str, "requirement": str, "method": str, "result": str}, ...]
        }
    """
    import docx

    path = _find_template_path(filename)
    doc = docx.Document(str(path))

    meta = {
        "tds": "",
        "cas": "",
        "opinion": "",
        "produkt_name": "",
    }

    # Extract product name from header table (col 1 typically contains
    # "ŚWIADECTWO JAKOŚCI\n/CERTIFICATE OF ANALYSIS\n<ProductName>")
    for section in doc.sections:
        hdr = section.header
        for tbl in hdr.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if "ŚWIADECTWO" in text.upper() or "CERTIFICATE" in text.upper():
                        # Last non-empty line after the title is product name
                        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
                        # Skip lines that are the title
                        for ln in lines:
                            if (
                                "ŚWIADECTWO" not in ln.upper()
                                and "CERTIFICATE" not in ln.upper()
                                and ln
                            ):
                                meta["produkt_name"] = ln
                                break

    # Extract from body paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # TDS and CAS on same line
        tds_match = re.search(r"\b([PT]\d{3,})\b", text)
        if tds_match:
            meta["tds"] = tds_match.group(1)

        cas_match = re.search(r"CAS:?\s*([\d\-]+)", text, re.IGNORECASE)
        if cas_match:
            meta["cas"] = cas_match.group(1)

        # Opinion line
        if "opinia" in text.lower() or "opinion" in text.lower():
            # The next paragraph(s) contain the opinion text — we capture here for now
            pass

        # Product compliance opinion (typically starts with "Produkt odpowiada")
        if "odpowiada" in text.lower() or "complies" in text.lower():
            meta["opinion"] = text

    # Extract table rows (first table = parameters table)
    header = []
    rows = []
    for table in doc.tables:
        if len(table.columns) < 3:
            continue
        for r_idx, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            if r_idx == 0:
                header = cells
                continue
            rows.append({
                "param": cells[0] if len(cells) > 0 else "",
                "requirement": cells[1] if len(cells) > 1 else "",
                "method": cells[2] if len(cells) > 2 else "",
                "result": cells[3] if len(cells) > 3 else "",
            })
        break  # only first multi-column table

    return {
        "meta": meta,
        "header": header,
        "rows": rows,
    }


def generate_certificate_pdf(filename: str, ebr: dict, wyniki_flat: dict) -> bytes:
    """Generate certificate PDF bytes.

    Args:
        filename: .docx template filename (basename only)
        ebr: dict with produkt, nr_partii, dt_start, status
        wyniki_flat: {kod: value} e.g. {"sm": 45.2, "nacl": 6.1}

    Returns:
        PDF bytes
    """
    from weasyprint import HTML

    parsed = parse_docx_template(filename)
    rows = parsed["rows"]

    # Fill results from CERT_MAPPINGS
    mappings = CERT_MAPPINGS.get(filename, [])
    for mapping in mappings:
        idx = mapping.get("row")
        kod = mapping.get("kod")
        if idx is None or idx >= len(rows):
            continue
        if kod and kod in wyniki_flat:
            val = wyniki_flat[kod]
            # Format: replace decimal point with comma for Polish notation
            if isinstance(val, float):
                rows[idx]["result"] = f"{val:.4g}".replace(".", ",")
            else:
                rows[idx]["result"] = str(val).replace(".", ",")

    # Fill dates
    today = date.today()
    dt_produkcji = ""
    dt_waznosci = ""
    if ebr.get("dt_start"):
        try:
            dt = datetime.fromisoformat(str(ebr["dt_start"]))
            dt_produkcji = dt.strftime("%d.%m.%Y")
            # +1 year
            dt_waznosci = dt.replace(year=dt.year + 1).strftime("%d.%m.%Y")
        except (ValueError, AttributeError):
            dt_produkcji = str(ebr["dt_start"])[:10]
            dt_waznosci = ""
    dt_wystawienia = today.strftime("%d.%m.%Y")

    html = render_template(
        "pdf/swiadectwo.html",
        meta=parsed["meta"],
        header=parsed["header"],
        rows=rows,
        ebr=ebr,
        dt_produkcji=dt_produkcji,
        dt_waznosci=dt_waznosci,
        dt_wystawienia=dt_wystawienia,
        filename=filename,
    )
    return HTML(string=html).write_pdf()


def save_certificate_pdf(
    pdf_bytes: bytes, produkt: str, template_filename: str, nr_partii: str
) -> str:
    """Save PDF to data/swiadectwa/{year}/{product}/{name}_{nr}.pdf.

    Returns:
        Relative path string (relative to project root).
    """
    year = date.today().year
    product_slug = produkt.replace(" ", "_")
    # Build a clean name from template filename
    template_base = template_filename
    if template_base.lower().endswith(".docx"):
        template_base = template_base[:-5]
    # Sanitize nr_partii for use in filename
    nr_safe = nr_partii.replace("/", "_").replace("\\", "_").replace(" ", "_")
    pdf_name = f"{template_base}_{nr_safe}.pdf"

    out_dir = OUTPUT_DIR / str(year) / product_slug
    out_dir.mkdir(parents=True, exist_ok=True)

    full_path = out_dir / pdf_name
    full_path.write_bytes(pdf_bytes)

    # Return relative path from project root
    project_root = Path(__file__).parent.parent
    return str(full_path.relative_to(project_root))
