"""Generate font-sample PDFs of a fake Świadectwo Jakości via Gotenberg.

Creates one DOCX per font combination, POSTs each to local Gotenberg
(http://localhost:3000), saves PDFs to ./font-samples/ for visual review.
"""
from __future__ import annotations

import io
from pathlib import Path

import requests
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


GOTENBERG_URL = "http://localhost:3000/forms/libreoffice/convert"
OUT_DIR = Path(__file__).resolve().parent.parent / "font-samples"


# Each variant defines fonts for: title, header, body, table, mono
VARIANTS = [
    # name, body_font, header_font, mono_font, sizes(title/h/body/tbl/mono), notes
    ("01_carlito_all",       "Carlito",          "Carlito",          "Liberation Mono",  (28, 13, 11, 10, 9), "Carlito everywhere — drop-in zamiennik Calibri"),
    ("02_caladea_carlito",   "Caladea",          "Carlito",          "Liberation Mono",  (28, 13, 11, 10, 9), "Caladea body (serif, klasyczny) + Carlito headers"),
    ("03_liberation_serif",  "Liberation Serif", "Liberation Sans",  "Liberation Mono",  (28, 13, 11, 10, 9), "Liberation Serif/Sans/Mono — Times+Arial+Courier feel"),
    ("04_dejavu_sans",       "DejaVu Sans",      "DejaVu Sans",      "DejaVu Sans Mono", (26, 13, 10, 9,  9), "DejaVu Sans całość — modern sans, świetny dla danych"),
    ("05_times_arial",       "Times New Roman",  "Arial",            "Courier New",     (28, 13, 11, 10, 9), "MS Core Fonts (Times+Arial+Courier) — klasyczny Word"),
    ("06_georgia_verdana",   "Georgia",          "Verdana",          "Liberation Mono", (26, 12, 10, 9,  9), "Georgia + Verdana — web-classic readability"),
    ("07_gentium_carlito",   "Gentium Basic",    "Carlito",          "DejaVu Sans Mono", (28, 13, 11, 10, 9), "Gentium body (akademicki serif) + Carlito headers"),
    ("08_noto",              "Noto Serif",       "Noto Sans",        "Noto Sans Mono",  (28, 13, 11, 10, 9), "Noto Serif/Sans/Mono — Google superrodzina"),
]


# Realistic-ish Świadectwo content with full Polish diacritics + sub/superscript
COMPANY = "P.W. iChegina Sp. z o.o."
ADDRESS = "ul. Przemysłowa 12, 41-200 Sosnowiec"
PRODUCT = "Płyn antykorozyjny K40GLO"
BATCH_NO = "K40GLO/2026/04/0287"
PROD_DATE = "2026-04-28"
EXP_DATE = "2027-04-28"

PARAMS = [
    # (name, unit, value, norm, method)
    ("Gęstość w 20°C",        "g/cm³",  "1,0823",     "1,080–1,085", "PN-EN ISO 3675"),
    ("Lepkość kinematyczna",  "mm²/s",  "12,4",       "11,5–13,0",   "PN-EN ISO 3104"),
    ("Liczba kwasowa",        "mg KOH/g","0,18",      "≤ 0,30",      "PN-ISO 6618"),
    ("Zawartość wody",        "% (m/m)", "0,032",     "≤ 0,050",     "PN-EN ISO 12937"),
    ("Temperatura zapłonu",   "°C",     "182",        "≥ 175",       "PN-EN ISO 2719"),
    ("Barwa wg skali Gardnera","—",     "2",          "≤ 3",         "PN-ISO 4630"),
    ("pH (10% wodny)",        "—",      "8,4",        "8,0–9,0",     "PN-EN 1262"),
    ("Klarowność",            "—",      "klarowny",   "klarowny",    "wizualna"),
]


def _shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _set_run(run, text, font, size_pt, *, bold=False, color=None):
    run.text = text
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), font)


def _para(doc, text, font, size_pt, *, bold=False, align=None, color=None, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run()
    _set_run(run, text, font, size_pt, bold=bold, color=color)
    return p


def _kv_row(doc, label, value, body_font, header_font, body_size, mono_font=None, mono_size=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run()
    _set_run(r1, f"{label}: ", header_font, body_size, bold=True)
    r2 = p.add_run()
    if mono_font:
        _set_run(r2, value, mono_font, mono_size or body_size, bold=False)
    else:
        _set_run(r2, value, body_font, body_size, bold=False)


def build_doc(variant) -> bytes:
    name, body_font, header_font, mono_font, sizes, _note = variant
    title_sz, h_sz, body_sz, tbl_sz, mono_sz = sizes

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # Header company line
    _para(doc, COMPANY, header_font, body_sz + 1, bold=True)
    _para(doc, ADDRESS, body_font, body_sz - 1, color="666666", space_after=10)

    # Title
    _para(doc, "ŚWIADECTWO JAKOŚCI", header_font, title_sz, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _para(doc, "Certificate of Analysis", body_font, body_sz, color="666666",
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

    # Identification block
    _kv_row(doc, "Wyrób", PRODUCT, body_font, header_font, body_sz)
    _kv_row(doc, "Numer partii", BATCH_NO, body_font, header_font, body_sz, mono_font=mono_font, mono_size=mono_sz + 1)
    _kv_row(doc, "Data produkcji", PROD_DATE, body_font, header_font, body_sz, mono_font=mono_font, mono_size=mono_sz + 1)
    _kv_row(doc, "Data ważności", EXP_DATE, body_font, header_font, body_sz, mono_font=mono_font, mono_size=mono_sz + 1)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Section header
    _para(doc, "Wyniki badań analitycznych", header_font, h_sz, bold=True, space_after=6)

    # Table: Parametr | Jednostka | Wartość | Norma | Metoda
    table = doc.add_table(rows=1, cols=5)
    table.autofit = False
    widths = [Cm(5.5), Cm(2.2), Cm(2.5), Cm(2.8), Cm(3.5)]
    headers = ["Parametr", "Jednostka", "Wartość", "Norma", "Metoda"]

    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.width = widths[i]
        _shading(cell, "E8E8E8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        _set_run(run, h, header_font, tbl_sz, bold=True)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for pname, unit, val, norm, method in PARAMS:
        row = table.add_row()
        for i, txt in enumerate([pname, unit, val, norm, method]):
            cell = row.cells[i]
            cell.width = widths[i]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run()
            # Numeric columns get mono font (for tabular alignment)
            use_mono = (i in (2,)) and mono_font
            font = mono_font if use_mono else body_font
            sz = mono_sz if use_mono else tbl_sz
            _set_run(run, txt, font, sz)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Conclusion
    _para(doc, "Wynik badań: zgodny z wymaganiami specyfikacji.",
          body_font, body_sz, bold=True, space_after=12)

    # Footer block (signatures)
    _para(doc, "Próbka pobrana zgodnie z procedurą PB-04.",
          body_font, body_sz - 1, color="666666", space_after=14)

    sig_table = doc.add_table(rows=2, cols=2)
    sig_cells = [
        ("Sporządził (laborant)", "mgr inż. Anna Kowalczyk"),
        ("Zatwierdził (technolog)", "dr inż. Łukasz Żółć"),
    ]
    for col, (label, person) in enumerate(sig_cells):
        c1 = sig_table.rows[0].cells[col]
        c2 = sig_table.rows[1].cells[col]
        c1_p = c1.paragraphs[0]
        _set_run(c1_p.add_run(), label, header_font, body_sz - 1, bold=True)
        c2_p = c2.paragraphs[0]
        _set_run(c2_p.add_run(), person, body_font, body_sz - 1)

    # Tiny footer with sample marker
    doc.add_paragraph()
    _para(doc, f"[Próbka czcionek: {name}]", mono_font, mono_sz - 1,
          color="999999", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def convert(docx_bytes: bytes, name: str) -> bytes:
    files = {"files": (f"{name}.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    r = requests.post(GOTENBERG_URL, files=files, timeout=60)
    r.raise_for_status()
    return r.content


def main():
    OUT_DIR.mkdir(exist_ok=True)
    print(f"→ output: {OUT_DIR}")
    for variant in VARIANTS:
        name, _, _, _, _, note = variant
        print(f"  • {name:30s}  {note}")
        docx_bytes = build_doc(variant)
        (OUT_DIR / f"{name}.docx").write_bytes(docx_bytes)
        try:
            pdf_bytes = convert(docx_bytes, name)
            (OUT_DIR / f"{name}.pdf").write_bytes(pdf_bytes)
        except Exception as exc:
            print(f"     !! Gotenberg fail: {exc}")
    print(f"✓ done — {len(VARIANTS)} variants in {OUT_DIR}")


if __name__ == "__main__":
    main()
